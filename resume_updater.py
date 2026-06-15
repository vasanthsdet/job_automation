"""
Updates a .docx resume with AI-tailored content for a specific job.

Strategy:
  - Copies base_resume.docx to tailored_resumes/<job_title>.docx
  - Finds the "Professional Summary" section → replaces with AI-generated summary
  - Finds the "Skills" section → replaces with AI-matched skills
  - All other content (experience, education, template formatting) is untouched
"""

import copy
import shutil
import os
import re
import json
from docx import Document
import anthropic
from config import ANTHROPIC_API_KEY

_client: anthropic.Anthropic | None = None

SUMMARY_KEYWORDS = {
    "professional summary", "summary", "profile",
    "objective", "about me", "professional profile", "career summary"
}
SKILLS_KEYWORDS = {
    "skills", "technical skills", "core competencies",
    "key skills", "expertise", "technologies", "tools & technologies"
}


def _client_instance() -> anthropic.Anthropic:
    global _client
    if _client is None:
        _client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    return _client


def _is_heading(para) -> bool:
    text = para.text.strip().lower()
    if not text:
        return False
    if "heading" in para.style.name.lower():
        return True
    # Bold short line = likely a section header
    has_bold = para.runs and all(r.bold for r in para.runs if r.text.strip())
    return has_bold and len(text) < 80


def _find_section(doc: Document, keywords: set) -> tuple[int, list[int]]:
    """Return (heading_idx, [content_para_idxs]) for the first section matching keywords."""
    paras = doc.paragraphs
    for i, para in enumerate(paras):
        if any(kw in para.text.strip().lower() for kw in keywords) and _is_heading(para):
            content = []
            j = i + 1
            while j < len(paras):
                if _is_heading(paras[j]) and paras[j].text.strip():
                    break
                if paras[j].text.strip():
                    content.append(j)
                j += 1
            return i, content
    return -1, []


def _set_para_text(para, new_text: str):
    """Replace paragraph text while preserving the first run's character formatting."""
    if not para.runs:
        para.add_run(new_text)
        return
    r = para.runs[0]
    saved = {
        "name": r.font.name,
        "size": r.font.size,
        "bold": r.bold,
        "italic": r.italic,
    }
    for run in para.runs:
        run.text = ""
    r.text = new_text
    r.font.name = saved["name"]
    r.font.size = saved["size"]
    r.bold = saved["bold"]
    r.italic = saved["italic"]


def _clone_para_after(ref_para, new_text: str):
    """Deep-clone ref_para, insert it after, set its text via XML."""
    from lxml import etree  # noqa: F401

    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    new_elem = copy.deepcopy(ref_para._element)
    ref_para._element.addnext(new_elem)
    t_elems = new_elem.findall(f".//{{{NS}}}t")
    if t_elems:
        t_elems[0].text = new_text
        for t in t_elems[1:]:
            t.text = ""


def get_resume_text(path: str) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _call_ai(job_title: str, job_description: str, resume_text: str) -> dict:
    prompt = f"""You are a professional QA resume writer specializing in CONTRACT roles.

JOB TITLE: {job_title}

JOB DESCRIPTION (first 3000 chars):
{job_description[:3000]}

CANDIDATE'S EXISTING RESUME (for context only):
{resume_text[:2000]}

Rules:
- This is a CONTRACT/hourly position — frame the summary to highlight contract readiness,
  quick ramp-up ability, and immediate availability.
- Only reference skills/tools visible in the existing resume or standard QA skills.
- Do NOT fabricate experience. Match JD keywords naturally.
- Emphasize tools, frameworks, and domain knowledge that match the JD.
- Return ONLY valid JSON, no markdown fences, no explanation.

JSON format:
{{
  "summary": "2-3 sentence ATS-friendly summary for this QA contract role, highlighting contract expertise and immediate value",
  "skills": ["Skill1", "Skill2", "Skill3", "Skill4", "Skill5", "Skill6", "Skill7", "Skill8"]
}}"""

    resp = _client_instance().messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=800,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = resp.content[0].text.strip()
    raw = re.sub(r"^```(?:json)?\n?", "", raw)
    raw = re.sub(r"\n?```$", "", raw)
    return json.loads(raw)


def create_tailored_resume(
    base_resume_path: str,
    job_title: str,
    job_description: str,
    output_dir: str = "tailored_resumes",
) -> str:
    """
    Returns path to the tailored .docx (or base_resume_path if tailoring fails).
    """
    os.makedirs(output_dir, exist_ok=True)
    safe = re.sub(r"[^\w\s-]", "", job_title).strip().replace(" ", "_")[:50]
    out_path = os.path.join(output_dir, f"Resume_{safe}.docx")
    shutil.copy2(base_resume_path, out_path)

    base_text = get_resume_text(base_resume_path)
    try:
        ai = _call_ai(job_title, job_description, base_text)
    except Exception as e:
        print(f"  [resume_updater] AI call failed: {e} — using base resume")
        return base_resume_path

    doc = Document(out_path)

    # --- Update summary ---
    _, s_idxs = _find_section(doc, SUMMARY_KEYWORDS)
    if s_idxs:
        _set_para_text(doc.paragraphs[s_idxs[0]], ai["summary"])
        for idx in s_idxs[1:]:
            _set_para_text(doc.paragraphs[idx], "")
    else:
        print("  [resume_updater] Summary section not found — skipping")

    # --- Update skills ---
    _, k_idxs = _find_section(doc, SKILLS_KEYWORDS)
    skills = ai.get("skills", [])
    if k_idxs and skills:
        paras = doc.paragraphs
        if len(k_idxs) == 1:
            # Single-line skills (pipe-separated)
            _set_para_text(paras[k_idxs[0]], " | ".join(skills))
        else:
            # Bullet-style skills
            for j, idx in enumerate(k_idxs):
                _set_para_text(paras[idx], skills[j] if j < len(skills) else "")
            # Append any extra skills by cloning the last bullet
            if len(skills) > len(k_idxs):
                ref = paras[k_idxs[-1]]
                for k in range(len(k_idxs), len(skills)):
                    _clone_para_after(ref, skills[k])
    else:
        print("  [resume_updater] Skills section not found — skipping")

    doc.save(out_path)
    print(f"  [resume_updater] Saved → {out_path}")
    return out_path
